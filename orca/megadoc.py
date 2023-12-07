"""TODO: File description."""
import json
import logging
from pathlib import Path
from typing import Iterable, Tuple

from docx import Document

log = logging.getLogger(__name__)


def get_headings(file_path: str | Path) -> Tuple[str, str]:
    """TODO: Description."""
    from dateutil.parser import parse

    # Try to use the format from icloud.py first. If that doesn't work, just
    # come back with the original filename so we at least have something.
    file_path = Path(file_path)
    try:
        parts = file_path.stem.split('_')

        date, time = parts[:2]
        timestamp = parse(f"{date} {time.replace('-', ':')}")
        timestamp = timestamp.strftime('%B %d, %Y at %I:%M %p')

        name = '_'.join(parts[2:])

        return name, timestamp

    except:
        return file_path.name, '[No timestamp.]'


def zip_files(file_paths: Iterable[str | Path], out_path: str | Path) -> None:
    """TODO: Description."""
    from zipfile import ZipFile

    log.info('Zipping files into %s...' % out_path)
    with ZipFile(Path(out_path).as_posix(), 'w') as zip:
        for file in [Path(p) for p in file_paths]:
            if file.exists() and file.is_file():
                zip.write(file, file.name)
    log.info('Done!')


def build_doc(data_path: str | Path, chunk_size: int = 5000) -> Path:
    """TODO: Description."""
    from natsort import natsorted
    from unidecode import unidecode

    data_path = Path(data_path)
    out_path = data_path / 'megadoc'
    out_path.mkdir(parents=True, exist_ok=True)

    json_files = natsorted(list(data_path.glob('*.json')))

    # Word docs can only realistically handle so much stuff at once before we
    # start hitting RAM limits. We can manage this by breaking each megadoc
    # into bite-sized chunks.
    file_count = len(json_files)
    chunk_total = file_count // chunk_size
    if file_count % chunk_size != 0:
        chunk_total += 1

    chunk_files = []
    for chunk in range(chunk_total):
        start = chunk * chunk_size
        end = (chunk + 1) * chunk_size

        doc = Document()
        filename = (
            f"{data_path.parent.name}_{(chunk + 1):02d}of{(chunk_total):02d}.docx"
        )
        doc_file = out_path / filename

        log.info('Building %s...' % doc_file)
        for i, json_file in enumerate(json_files[start:end]):
            log.info('%s (%d/%d)' % (json_file, i + start + 1, file_count))
            with json_file.open() as f:
                data = json.load(f)

            # Get heading from filename.
            name, timestamp = get_headings(json_file)
            doc.add_heading(name, level=1)
            doc.add_heading(timestamp, level=2)

            # Different Azure models return data in different formats. We can
            # tell which one was used by the way key used to store the result.
            # TODO: Break this out into separate function definition(s).
            # ... Computer Vision
            if 'readResult' in data:
                block_key = 'blocks' if 'blocks' in data['readResult'] else 'pages'
                for block in data['readResult'][block_key]:
                    lines = []
                    for line in block['lines']:
                        text_key = 'text' if 'text' in line else 'content'
                        lines.append(unidecode(line[text_key]))
                    doc.add_paragraph('\n'.join(lines))

            # ... Document Intelligence
            elif 'analyzeResult' in data:
                try:
                    for p in data['analyzeResult']['paragraphs']:
                        doc.add_paragraph(unidecode(p['content']))
                except KeyError:
                    doc.add_paragraph('[No text recovered.]')

            # ... something went wrong; OCR errored out somehow.
            else:
                doc.add_paragraph('[No text recovered.]')

            doc.add_page_break()

        log.info('Saving %s...' % doc_file)
        doc.save(doc_file.as_posix())
        chunk_files.append(doc_file)

    zip_files(chunk_files, out_path / f"{data_path.parent.name}_{data_path.name}.zip")
    log.info('Done!')


if __name__ == '__main__':
    import argparse
    from dotenv import load_dotenv

    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
    )

    load_dotenv()
    parser = argparse.ArgumentParser()
    parser.add_argument(
        'paths',
        metavar='path',
        type=str,
        nargs='+',
        help='paths containing json files',
    )
    args = parser.parse_args()

    # Kludge: Sometimes the name of the file gets scooped up in here.
    paths = [p for p in args.paths if 'megadoc.py' not in p]
    for path in [Path(p) for p in paths]:
        build_doc(path)
